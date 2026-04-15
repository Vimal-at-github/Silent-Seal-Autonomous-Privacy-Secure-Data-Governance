"""
SilentSeal - File Watcher (Standby Mode)
Real-time file system monitoring using watchdog
"""

import os
import threading
import time
from pathlib import Path
from typing import List, Dict, Any, Callable, Set
from dataclasses import dataclass
from enum import Enum


class WatcherStatus(Enum):
    """Watcher status states"""
    STOPPED = "stopped"
    RUNNING = "running"
    PAUSED = "paused"


@dataclass
class WatchEvent:
    """Represents a file system event"""
    event_type: str  # created, modified, deleted, moved
    file_path: str
    is_directory: bool
    timestamp: float


class FileWatcher:
    """
    Real-time file system watcher for Standby Mode.
    
    Features:
    - Monitor Downloads/Documents/custom folders
    - Detect new files automatically
    - Trigger PII scans on file creation
    - Integration with notification system
    - Start/Stop/Pause controls
    """
    
    # Supported file extensions for scanning
    SUPPORTED_EXTENSIONS = {
        # Documents
        '.pdf', '.doc', '.docx', '.odt', '.rtf',
        # Spreadsheets
        '.xls', '.xlsx', '.ods', '.csv',
        # Text files
        '.txt', '.md', '.log', '.json', '.xml', '.html',
        # Images (for OCR in future)
        '.jpg', '.jpeg', '.png', '.tiff', '.bmp'
    }
    
    def __init__(self, 
                 on_new_file: Callable[[str], None] = None,
                 on_sensitive_detected: Callable[[str, Dict], None] = None):
        """
        Initialize the file watcher.
        
        Args:
            on_new_file: Callback when a new file is detected
            on_sensitive_detected: Callback when sensitive data is found
        """
        self.status = WatcherStatus.STOPPED
        self.watched_paths: List[str] = []
        self.observers: List[Any] = []  # Initialize observers list
        # Event handling
        self._lock = threading.Lock()
        self._event_queue: List[WatchEvent] = []
        self._processed_files: set = set()
        self._debounce_delay = 2.0  # seconds
        
        # Recent detections for UI
        self.recent_detections: List[Dict[str, Any]] = []
        self.max_recent = 50  # Keep last 50 detections
        self._persistence_file = os.path.join(os.path.dirname(__file__), "..", "database", "recent_detections.json")
        self._load_detections()
        
        # Watcher statecks
        self.on_new_file = on_new_file
        self.on_sensitive_detected = on_sensitive_detected
        
        # Processing thread
        self._processor_thread = None
        self._stop_event = threading.Event()
        
        # Default paths to watch
        self.default_paths = self._get_default_watch_paths()
    
    def _load_detections(self):
        """Load recent detections from JSON file"""
        try:
            if os.path.exists(self._persistence_file):
                import json
                with open(self._persistence_file, 'r') as f:
                    self.recent_detections = json.load(f)
                print(f"📊 Loaded {len(self.recent_detections)} detections from persistence")
        except Exception as e:
            print(f"⚠️ Failed to load detections: {e}")

    def _save_detections(self):
        """Save recent detections to JSON file"""
        try:
            import json
            os.makedirs(os.path.dirname(self._persistence_file), exist_ok=True)
            with open(self._persistence_file, 'w') as f:
                json.dump(self.recent_detections, f, indent=2)
        except Exception as e:
            print(f"⚠️ Failed to save detections: {e}")
    
    def _get_default_watch_paths(self) -> List[str]:
        """Get default paths to watch (Downloads and Documents)"""
        home = Path.home()
        paths = []
        
        # Windows paths
        downloads = home / "Downloads"
        documents = home / "Documents"
        
        if downloads.exists():
            paths.append(str(downloads))
        if documents.exists():
            paths.append(str(documents))
        
        return paths
    
    def start(self, paths: List[str] = None) -> Dict[str, Any]:
        """
        Start watching specified paths.
        
        Args:
            paths: List of directory paths to watch. Defaults to Downloads/Documents.
            
        Returns:
            Status dict with watched paths
        """
        if self.status == WatcherStatus.RUNNING:
            return {"status": "already_running", "paths": self.watched_paths}
        
        try:
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler
        except ImportError:
            return {"status": "error", "message": "watchdog library not installed"}
        
        # Use default paths if none specified
        paths = paths or self.default_paths
        
        # Validate paths
        valid_paths = []
        for path in paths:
            if os.path.exists(path) and os.path.isdir(path):
                valid_paths.append(path)
        
        if not valid_paths:
            return {"status": "error", "message": "No valid paths to watch"}
        
        self.watched_paths = valid_paths
        
        # Create event handler
        handler = self._create_event_handler()
        
        # Create observers for each path
        self.observers = []
        for path in valid_paths:
            observer = Observer()
            observer.schedule(handler, path, recursive=True)
            observer.start()
            self.observers.append(observer)
        
        # Start processor thread
        self._stop_event.clear()
        self._processor_thread = threading.Thread(
            target=self._process_events, 
            daemon=True
        )
        self._processor_thread.start()
        
        self.status = WatcherStatus.RUNNING
        
        return {
            "status": "started",
            "paths": self.watched_paths,
            "extensions": list(self.SUPPORTED_EXTENSIONS)
        }
    
    def stop(self) -> Dict[str, Any]:
        """Stop all file watching"""
        if self.status == WatcherStatus.STOPPED:
            return {"status": "already_stopped"}
        
        # Stop observers
        for observer in self.observers:
            observer.stop()
            observer.join(timeout=2)
        
        self.observers = []
        
        # Stop processor thread
        self._stop_event.set()
        if self._processor_thread:
            self._processor_thread.join(timeout=2)
        
        self.status = WatcherStatus.STOPPED
        
        return {"status": "stopped", "files_processed": len(self._processed_files)}
    
    def pause(self) -> Dict[str, Any]:
        """Pause watching (keeps observers but doesn't process events)"""
        if self.status != WatcherStatus.RUNNING:
            return {"status": "not_running"}
        
        self.status = WatcherStatus.PAUSED
        return {"status": "paused"}
    
    def resume(self) -> Dict[str, Any]:
        """Resume watching after pause"""
        if self.status != WatcherStatus.PAUSED:
            return {"status": "not_paused"}
        
        self.status = WatcherStatus.RUNNING
        return {"status": "resumed"}
    
    def get_status(self) -> Dict[str, Any]:
        """Get current watcher status"""
        return {
            "status": self.status.value,
            "watched_paths": self.watched_paths,
            "files_processed": len(self._processed_files),
            "pending_events": len(self._event_queue)
        }
    
    def _create_event_handler(self):
        """Create the watchdog event handler"""
        from watchdog.events import FileSystemEventHandler
        
        watcher = self  # Reference for inner class
        
        class SilentSealHandler(FileSystemEventHandler):
            def on_created(self, event):
                if not event.is_directory:
                    watcher._handle_file_event("created", event.src_path)
            
            def on_modified(self, event):
                # Ignore modified events to prevent duplicate processing
                # Files are scanned on creation, not modification
                pass
            
            def on_moved(self, event):
                if not event.is_directory:
                    watcher._handle_file_event("moved", event.dest_path)
        
        return SilentSealHandler()
    
    def _handle_file_event(self, event_type: str, file_path: str):
        """Handle a file system event"""
        # Check file extension
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return
        
        print(f"📁 File event detected: {event_type} - {Path(file_path).name}")
        
        # Skip if already processed recently
        if file_path in self._processed_files:
            print(f"⏭️  Skipping already processed: {Path(file_path).name}")
            return
        
        # Add to event queue
        with self._lock:
            self._event_queue.append(WatchEvent(
                event_type=event_type,
                file_path=file_path,
                is_directory=False,
                timestamp=time.time()
            ))
            print(f"✓ Added to queue: {Path(file_path).name}")
    
    
    def _process_events(self):
        """Background thread to process file events with debouncing"""
        while not self._stop_event.is_set():
            if self.status != WatcherStatus.RUNNING:
                time.sleep(0.5)
                continue
            
            # Get events that are ready (past debounce delay)
            ready_events = []
            current_time = time.time()
            
            with self._lock:
                remaining_events = []
                for event in self._event_queue:
                    if current_time - event.timestamp >= self._debounce_delay:
                        ready_events.append(event)
                    else:
                        remaining_events.append(event)
                self._event_queue = remaining_events
            
            # Process ready events
            for event in ready_events:
                self._process_single_event(event)
            
            time.sleep(0.5)
    
    def _process_single_event(self, event: WatchEvent):
        """Process a single file event"""
        file_path = event.file_path
        
        # Mark as processed
        self._processed_files.add(file_path)
        
        # Limit processed files cache size
        if len(self._processed_files) > 10000:
            self._processed_files = set(list(self._processed_files)[-5000:])
        
        # Callback for new file
        if self.on_new_file:
            try:
                self.on_new_file(file_path)
            except Exception as e:
                print(f"Error in on_new_file callback: {e}")
        
        # Scan for sensitive data
        self._scan_file_for_pii(file_path)
    
    def _scan_file_for_pii(self, file_path: str):
        """Scan a file for PII and trigger notifications"""
        try:
            # Check if file still exists and is accessible
            if not os.path.exists(file_path):
                return
            
            # Try lightweight text extraction first
            text = self._extract_text_simple(file_path)
            if not text:
                print(f"⚠️  No text extracted from {os.path.basename(file_path)}")
                return
            
            # ---------------------------------------------------------
            # Deduplication: Check if content hash matches recently processed
            # ---------------------------------------------------------
            import hashlib
            content_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
            
            # Simple in-memory cache for hashes (keep last 100)
            if not hasattr(self, '_processed_hashes'):
                self._processed_hashes = []
            
            if content_hash in self._processed_hashes:
                print(f"⏭️  Duplicate content detected for {os.path.basename(file_path)}, skipping.")
                return

            self._processed_hashes.append(content_hash)
            if len(self._processed_hashes) > 100:
                self._processed_hashes.pop(0)
            # ---------------------------------------------------------
            
            print(f"📄 Extracted {len(text)} characters from {os.path.basename(file_path)}")
            
            # Use lightweight regex detection
            entities = self._detect_with_regex(text)
            
            print(f"🔍 Found {len(entities)} entities in {os.path.basename(file_path)}")
            
            if not entities or len(entities) == 0:
                print(f"⏭️  No PII found, skipping {os.path.basename(file_path)}")
                return  # No PII found
            
            # Calculate simple risk score
            risk_score = min(len(entities) * 15, 100)
            if risk_score > 70:
                risk_level = "HIGH"
            elif risk_score > 40:
                risk_level = "MEDIUM"
            else:
                risk_level = "LOW"
            
            # Store detection for UI
            detection_info = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "timestamp": time.time(),
                "risk_level": risk_level,
                "risk_score": risk_score,
                "entities_count": len(entities),
                "entities": [{"type": e["type"], "value": e["value"][:20] + "..." if len(e["value"]) > 20 else e["value"]} for e in entities[:10]]  # Truncate for privacy
            }
            
            with self._lock:
                self.recent_detections.insert(0, detection_info)  # Add to front
                if len(self.recent_detections) > self.max_recent:
                    self.recent_detections = self.recent_detections[:self.max_recent]
                print(f"📊 Added to recent_detections. Total count: {len(self.recent_detections)}")
                self._save_detections()
            
            # Send notification
            try:
                from features.notifications import get_notification_manager
                notifier = get_notification_manager()
                result = notifier.notify_sensitive_file(
                    file_path=file_path,
                    risk_level=risk_level,
                    entities_count=len(entities),
                    risk_score=risk_score
                )
                if not result:
                    print(f"⚠️  Desktop notification failed for {os.path.basename(file_path)}")
                else:
                    print(f"✓ Desktop notification sent for {os.path.basename(file_path)}")
            except Exception as e:
                print(f"Notification error: {e}")
                import traceback
                traceback.print_exc()
            
            # Callback with results
            if self.on_sensitive_detected:
                self.on_sensitive_detected(file_path, {
                    "entities": entities,
                    "risk": {"score": risk_score, "level": risk_level}
                })
            
            print(f"✓ Detected {len(entities)} PII entities in {Path(file_path).name}")
                
        except Exception as e:
            print(f"Error scanning file {file_path}: {e}")
    
    def _extract_text_simple(self, file_path: str) -> str:
        """Comprehensive text extraction supporting multiple formats"""
        try:
            ext = Path(file_path).suffix.lower()
            
            # Plain text files
            if ext in ['.txt', '.csv', '.json', '.xml', '.html', '.log', '.md']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read(100000)  # First 100KB
                except:
                    # Try with different encoding
                    with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
                        return f.read(100000)
            
            # PDF files
            if ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        # Extract from first 20 pages
                        for i, page in enumerate(reader.pages[:20]):
                            if i >= 20:
                                break
                            try:
                                text += page.extract_text() + "\n"
                            except:
                                continue
                        return text
                except Exception as e:
                    print(f"PDF extraction error: {e}")
                    return ""
            
            # Word documents (.docx)
            if ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = []
                    for para in doc.paragraphs[:500]:  # First 500 paragraphs
                        text.append(para.text)
                    # Also extract from tables
                    for table in doc.tables[:20]:  # First 20 tables
                        for row in table.rows:
                            for cell in row.cells:
                                text.append(cell.text)
                    return "\n".join(text)
                except Exception as e:
                    print(f"DOCX extraction error: {e}")
                    return ""
            
            # Excel files (.xlsx, .xls)
            if ext in ['.xlsx', '.xls']:
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                    text = []
                    # Read first 5 sheets
                    for sheet_name in list(wb.sheetnames)[:5]:
                        ws = wb[sheet_name]
                        # Read first 1000 rows
                        for row in list(ws.iter_rows(max_row=1000, values_only=True)):
                            row_text = " ".join([str(cell) for cell in row if cell is not None])
                            if row_text.strip():
                                text.append(row_text)
                    return "\n".join(text)
                except Exception as e:
                    print(f"Excel extraction error: {e}")
                    return ""
            
            # RTF files
            if ext == '.rtf':
                try:
                    # Simple RTF text extraction (strips formatting)
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(100000)
                        # Basic RTF cleanup - remove control words
                        import re
                        # Remove RTF control sequences
                        text = re.sub(r'\\[a-z]+\d*\s?', ' ', content)
                        text = re.sub(r'[{}]', '', text)
                        return text
                except Exception as e:
                    print(f"RTF extraction error: {e}")
                    return ""
            
            # For unsupported formats, return empty
            print(f"Unsupported format: {ext}")
            return ""
                
        except Exception as e:
            print(f"Text extraction error for {file_path}: {e}")
            return ""
    
    def _detect_with_regex(self, text: str) -> list:
        """Lightweight regex-based entity detection"""
        import re
        
        entities = []
        
        # Email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for match in re.finditer(email_pattern, text):
            entities.append({
                "type": "EMAIL", 
                "value": match.group(),
                "position": match.span()
            })
        
        # Phone numbers (various formats)
        phone_patterns = [
            r'\b\d{10}\b',  # 10 digits
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # XXX-XXX-XXXX
            r'\+\d{1,3}[-.\s]?\d{10}\b',  # International
        ]
        for pattern in phone_patterns:
            for match in re.finditer(pattern, text):
                entities.append({
                    "type": "PHONE",
                    "value": match.group(),
                    "position": match.span()
                })
        
        # Aadhaar numbers (12 digits)
        aadhaar_pattern = r'\b\d{4}\s?\d{4}\s?\d{4}\b'
        for match in re.finditer(aadhaar_pattern, text):
            entities.append({
                "type": "AADHAAR",
                "value": match.group(),
                "position": match.span()
            })
        
        # PAN card (AAAAA9999A format)
        pan_pattern = r'\b[A-Z]{5}[0-9]{4}[A-Z]\b'
        for match in re.finditer(pan_pattern, text):
            entities.append({
                "type": "PAN",
                "value": match.group(),
                "position": match.span()
            })
        
        # Credit card numbers (13-16 digits with optional dashes/spaces)
        cc_pattern = r'\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{3,4}\b'
        for match in re.finditer(cc_pattern, text):
            entities.append({
                "type": "CREDIT_CARD",
                "value": match.group(),
                "position": match.span()
            })
        
        # SSN and similar formats
        ssn_pattern = r'\b\d{3}-\d{2}-\d{4}\b'
        for match in re.finditer(ssn_pattern, text):
            entities.append({
                "type": "SSN",
                "value": match.group(),
                "position": match.span()
            })
        
        return entities
    
    def add_watch_path(self, path: str) -> Dict[str, Any]:
        """Add a new path to watch"""
        if not os.path.exists(path) or not os.path.isdir(path):
            return {"status": "error", "message": "Invalid path"}
        
        if path in self.watched_paths:
            return {"status": "already_watching", "path": path}
        
        try:
            from watchdog.observers import Observer
            
            handler = self._create_event_handler()
            observer = Observer()
            observer.schedule(handler, path, recursive=True)
            
            if self.status == WatcherStatus.RUNNING:
                observer.start()
            
            self.observers.append(observer)
            self.watched_paths.append(path)
            
            return {"status": "added", "path": path}
            
        except Exception as e:
            return {"status": "error", "message": str(e)}
    
    def remove_watch_path(self, path: str) -> Dict[str, Any]:
        """Remove a path from watching"""
        if path not in self.watched_paths:
            return {"status": "not_watching", "path": path}
        
        idx = self.watched_paths.index(path)
        if idx < len(self.observers):
            observer = self.observers[idx]
            observer.stop()
            observer.join(timeout=2)
            self.observers.pop(idx)
        
        self.watched_paths.remove(path)
        
        return {"status": "removed", "path": path}


# Global watcher instance
_file_watcher = None

def get_file_watcher() -> FileWatcher:
    """Get or create the global file watcher"""
    global _file_watcher
    if _file_watcher is None:
        _file_watcher = FileWatcher()
    return _file_watcher
